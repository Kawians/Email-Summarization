import os
import email
from email import policy
from email.parser import BytesParser
from docx import Document
from nltk.tokenize import sent_tokenize
from nltk.corpus import stopwords
from nltk.probability import FreqDist
from nltk.tokenize import word_tokenize
import nltk

nltk.download('punkt')
nltk.download('stopwords')

def extract_text_from_eml(eml_path):
    with open(eml_path, 'rb') as file:
        msg = BytesParser(policy=policy.default).parse(file)
        return msg.get_body(preferencelist=('plain', 'html')).get_content()

def summarize_text(text, summary_length=5):
    sentences = sent_tokenize(text)
    words = word_tokenize(text)
    words = [word.lower() for word in words if word.isalpha()]
    stop_words = set(stopwords.words('english'))
    words = [word for word in words if word not in stop_words]
    
    freq_dist = FreqDist(words)
    ranked_sentences = sorted(sentences, key=lambda sent: sum(freq_dist[word] for word in word_tokenize(sent) if word in freq_dist), reverse=True)
    
    return ' '.join(ranked_sentences[:summary_length])

def combine_eml_to_word(folder_path, output_combined, output_summary):
    combined_text = ""
    for filename in sorted(os.listdir(folder_path)):
        if filename.endswith('.eml'):
            eml_path = os.path.join(folder_path, filename)
            text = extract_text_from_eml(eml_path)
            combined_text += text + "\n\n"

    # Create combined Word document
    combined_doc = Document()
    combined_doc.add_heading('Combined Email Conversations', 0)
    combined_doc.add_paragraph(combined_text)
    combined_doc.save(output_combined)
    
    # Create summary Word document
    summary = summarize_text(combined_text)
    summary_doc = Document()
    summary_doc.add_heading('Summary of Email Conversations', 0)
    summary_doc.add_paragraph(summary)
    summary_doc.save(output_summary)

folder_path = '/Users/kavian/Desktop/Scotia/Email Archive'
output_combined = '/Users/kavian/Desktop/Scotia/combined_emails.docx'
output_summary = '/Users/kavian/Desktop/Scotia/summary_emails.docx'

combine_eml_to_word(folder_path, output_combined, output_summary)
